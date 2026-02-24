"""Tests for cwac_mcp.report_generator."""

import os
from datetime import datetime
from unittest.mock import patch

import pytest


class TestGenerateMarkdownReport:
    """Tests for markdown report generation."""

    def test_renders_scan_summary(self, sample_report_context):
        from cwac_mcp.report_generator import generate_markdown_report

        md = generate_markdown_report("cwac_scan_report", sample_report_context)
        assert "test_scan" in md
        assert "example.govt.nz" in md

    def test_includes_issue_counts(self, sample_report_context):
        from cwac_mcp.report_generator import generate_markdown_report

        md = generate_markdown_report("cwac_scan_report", sample_report_context)
        assert "critical" in md.lower()
        assert "serious" in md.lower()

    def test_renders_summary_template(self, sample_scan_summary):
        from cwac_mcp.report_generator import generate_markdown_report

        context = {
            "audit_name": "summary_test",
            "scan_date": "2026-02-24",
            "summary": sample_scan_summary,
            "generated_at": datetime.now().isoformat(),
        }
        md = generate_markdown_report("cwac_summary_report", context)
        assert "summary_test" in md

    def test_renders_visual_template(self, sample_visual_findings):
        from cwac_mcp.report_generator import generate_markdown_report

        context = {
            "url": "https://example.com/team/",
            "scan_date": "2026-02-24",
            "findings": sample_visual_findings,
            "total_findings": len(sample_visual_findings),
            "generated_at": datetime.now().isoformat(),
        }
        md = generate_markdown_report("visual_scan_report", context)
        assert "Heading-like content" in md

    def test_unknown_template_raises(self):
        from cwac_mcp.report_generator import generate_markdown_report

        with pytest.raises(FileNotFoundError):
            generate_markdown_report("nonexistent_template", {})


class TestGenerateDocxReport:
    """Tests for DOCX report generation."""

    def test_creates_docx_file(self, sample_report_context, tmp_output_dir):
        from cwac_mcp.report_generator import generate_docx_report

        path = os.path.join(tmp_output_dir, "test_report.docx")
        generate_docx_report("cwac_scan_report", sample_report_context, path)
        assert os.path.isfile(path)
        assert os.path.getsize(path) > 0

    def test_docx_is_valid(self, sample_report_context, tmp_output_dir):
        from docx import Document

        from cwac_mcp.report_generator import generate_docx_report

        path = os.path.join(tmp_output_dir, "test_report.docx")
        generate_docx_report("cwac_scan_report", sample_report_context, path)
        doc = Document(path)
        assert len(doc.paragraphs) > 0

    def test_docx_contains_audit_name(self, sample_report_context, tmp_output_dir):
        from docx import Document

        from cwac_mcp.report_generator import generate_docx_report

        path = os.path.join(tmp_output_dir, "test_report.docx")
        generate_docx_report("cwac_scan_report", sample_report_context, path)
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "test_scan" in text


class TestGenerateReports:
    """Tests for the combined report generation function."""

    def test_generates_both_formats(self, sample_report_context, tmp_output_dir):
        from cwac_mcp.report_generator import generate_reports

        paths = generate_reports(
            template_name="cwac_scan_report",
            context=sample_report_context,
            output_dir=tmp_output_dir,
            audit_name="test_scan",
        )
        assert "md" in paths
        assert "docx" in paths
        assert os.path.isfile(paths["md"])
        assert os.path.isfile(paths["docx"])

    def test_output_filenames_contain_audit_name(self, sample_report_context, tmp_output_dir):
        from cwac_mcp.report_generator import generate_reports

        paths = generate_reports(
            template_name="cwac_scan_report",
            context=sample_report_context,
            output_dir=tmp_output_dir,
            audit_name="my_audit",
        )
        assert "my_audit" in os.path.basename(paths["md"])
        assert "my_audit" in os.path.basename(paths["docx"])

    def test_output_filenames_contain_timestamp(self, sample_report_context, tmp_output_dir):
        from cwac_mcp.report_generator import generate_reports

        paths = generate_reports(
            template_name="cwac_scan_report",
            context=sample_report_context,
            output_dir=tmp_output_dir,
            audit_name="test_scan",
        )
        # Filename should contain a date-like pattern
        basename = os.path.basename(paths["md"])
        assert "202" in basename  # year prefix

    def test_creates_output_dir_if_missing(self, sample_report_context, tmp_path):
        from cwac_mcp.report_generator import generate_reports

        output_dir = str(tmp_path / "new_output")
        paths = generate_reports(
            template_name="cwac_scan_report",
            context=sample_report_context,
            output_dir=output_dir,
            audit_name="test_scan",
        )
        assert os.path.isdir(output_dir)
        assert os.path.isfile(paths["md"])
