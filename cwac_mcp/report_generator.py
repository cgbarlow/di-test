"""Report generator for CWAC and visual scan results.

Generates accessibility scan reports in Markdown (via Jinja2 templates) and
DOCX (via python-docx) formats. Reports auto-save to ./output/ with
filenames: {audit_name}_{timestamp}_report.{md,docx}
"""

import os
from datetime import datetime

from jinja2 import Environment, FileSystemLoader

# Template directory relative to this file's parent (project root).
_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_TEMPLATE_DIR = os.path.join(_PROJECT_ROOT, "templates")


def _get_jinja_env() -> Environment:
    """Create a Jinja2 environment configured for report templates."""
    return Environment(
        loader=FileSystemLoader(_TEMPLATE_DIR),
        trim_blocks=True,
        lstrip_blocks=True,
        keep_trailing_newline=True,
    )


def generate_markdown_report(template_name: str, context: dict) -> str:
    """Render a markdown report from a Jinja2 template.

    Args:
        template_name: Template name without extension (e.g. "cwac_scan_report").
            The ".md.j2" extension is appended automatically.
        context: Template context dict with report data.

    Returns:
        The rendered markdown string.

    Raises:
        FileNotFoundError: If the template file does not exist.
    """
    env = _get_jinja_env()
    template_file = f"{template_name}.md.j2"

    try:
        template = env.get_template(template_file)
    except Exception:
        template_path = os.path.join(_TEMPLATE_DIR, template_file)
        raise FileNotFoundError(f"Template not found: {template_path}")

    return template.render(**context)


def generate_docx_report(template_name: str, context: dict, output_path: str) -> None:
    """Generate a DOCX report from structured data.

    Uses python-docx to build the document directly from the context data,
    without requiring pandoc.

    Args:
        template_name: Template name (used to select the builder).
        context: Template context dict with report data.
        output_path: Absolute path where the .docx file will be saved.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    audit_name = context.get("audit_name", context.get("url", "Accessibility Report"))
    title = doc.add_heading(f"Accessibility Report: {audit_name}", level=0)

    # Metadata
    scan_date = context.get("scan_date", context.get("generated_at", ""))
    if scan_date:
        doc.add_paragraph(f"Generated: {scan_date}")

    if template_name == "cwac_scan_report":
        _build_cwac_scan_docx(doc, context)
    elif template_name == "cwac_summary_report":
        _build_cwac_summary_docx(doc, context)
    elif template_name == "visual_scan_report":
        _build_visual_scan_docx(doc, context)

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def _build_cwac_scan_docx(doc, context: dict) -> None:
    """Build DOCX content for a CWAC scan report."""
    from docx.shared import Pt

    summary = context.get("summary", {})

    # Summary section
    doc.add_heading("Summary", level=1)

    total = context.get("total_issues", summary.get("total_issues", 0))
    pages = context.get("pages_scanned", "N/A")
    doc.add_paragraph(f"Total issues: {total}")
    doc.add_paragraph(f"Pages scanned: {pages}")

    # Impact breakdown
    impact = summary.get("axe_impact_breakdown", {})
    if impact:
        doc.add_heading("Impact Breakdown", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Impact"
        hdr[1].text = "Count"
        for level, count in impact.items():
            row = table.add_row().cells
            row[0].text = str(level)
            row[1].text = str(count)

    # Top violations
    violations = summary.get("top_violations", [])
    if violations:
        doc.add_heading("Top Violations", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Rule"
        hdr[1].text = "Count"
        for v in violations:
            row = table.add_row().cells
            row[0].text = str(v.get("id", ""))
            row[1].text = str(v.get("count", 0))

    # Detailed findings
    results = context.get("results", [])
    if results:
        doc.add_heading("Detailed Findings", level=1)
        for i, result in enumerate(results, 1):
            doc.add_heading(f"Finding {i}: {result.get('id', 'Unknown')}", level=2)
            doc.add_paragraph(f"Impact: {result.get('impact', 'N/A')}")
            doc.add_paragraph(f"Description: {result.get('description', 'N/A')}")
            doc.add_paragraph(f"URL: {result.get('url', 'N/A')}")
            if result.get("html"):
                doc.add_paragraph(f"HTML: {result['html']}")


def _build_cwac_summary_docx(doc, context: dict) -> None:
    """Build DOCX content for a CWAC summary report."""
    summary = context.get("summary", {})

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(f"Total issues: {summary.get('total_issues', 0)}")

    by_type = summary.get("by_audit_type", {})
    if by_type:
        doc.add_heading("Issues by Audit Type", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Audit Type"
        hdr[1].text = "Count"
        for audit_type, count in by_type.items():
            row = table.add_row().cells
            row[0].text = str(audit_type)
            row[1].text = str(count)


def _build_visual_scan_docx(doc, context: dict) -> None:
    """Build DOCX content for a visual scan report."""
    doc.add_heading("Visual Pattern Findings", level=1)

    url = context.get("url", "N/A")
    doc.add_paragraph(f"Scanned URL: {url}")
    doc.add_paragraph(f"Total findings: {context.get('total_findings', 0)}")

    findings = context.get("findings", [])
    for i, finding in enumerate(findings, 1):
        doc.add_heading(f"Finding {i}: {finding.get('type', 'Unknown')}", level=2)
        doc.add_paragraph(f"Reason: {finding.get('reason', 'N/A')}")

        location = finding.get("location", {})
        if location:
            doc.add_paragraph(f"CSS Selector: {location.get('cssSelector', 'N/A')}")

        visual = finding.get("visual", {})
        if visual:
            doc.add_paragraph(f"Font size: {visual.get('fontSize', 'N/A')}")
            doc.add_paragraph(f"Font weight: {visual.get('fontWeight', 'N/A')}")

        if finding.get("htmlSnippet"):
            doc.add_paragraph(f"HTML: {finding['htmlSnippet']}")

        confidence = finding.get("confidence")
        if confidence is not None:
            doc.add_paragraph(f"Confidence: {confidence}")


def _build_output_filename(audit_name: str, extension: str) -> str:
    """Build a report output filename with timestamp.

    Args:
        audit_name: The audit name to include in the filename.
        extension: File extension without dot (e.g. "md", "docx").

    Returns:
        Filename string: {audit_name}_{timestamp}_report.{extension}
    """
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    # Sanitize audit_name for filename
    import re
    safe_name = re.sub(r"[^a-zA-Z0-9_\-.]", "_", audit_name)
    return f"{safe_name}_{timestamp}_report.{extension}"


def generate_reports(
    template_name: str,
    context: dict,
    output_dir: str,
    audit_name: str,
) -> dict[str, str]:
    """Generate reports in both Markdown and DOCX formats.

    Args:
        template_name: Template name (e.g. "cwac_scan_report").
        context: Template context dict with report data.
        output_dir: Directory to save reports in.
        audit_name: Audit name for the output filename.

    Returns:
        Dict mapping format to output file path: {"md": path, "docx": path}.
    """
    os.makedirs(output_dir, exist_ok=True)

    # Generate markdown
    md_content = generate_markdown_report(template_name, context)
    md_filename = _build_output_filename(audit_name, "md")
    md_path = os.path.join(output_dir, md_filename)
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    # Generate DOCX
    docx_filename = _build_output_filename(audit_name, "docx")
    docx_path = os.path.join(output_dir, docx_filename)
    generate_docx_report(template_name, context, docx_path)

    return {"md": md_path, "docx": docx_path}
